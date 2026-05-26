"""生成一个带故意错字的测试 docx,用于 spike 验证。"""
from pathlib import Path
from docx import Document


def make_test_docx(path: Path):
    doc = Document()
    doc.add_heading("测试文档:乡里旧事", level=1)

    paragraphs = [
        "他听到这话,心里一颤,真有种心惊胆颤的感觉。",
        "村子背后的山青水秀,远远望去像一幅画。",
        "在其时,他还是个孩子,根本不懂大人的世界。",
        "既使下着雨,他也要走完这段路。",
        "会议对当前的问题进行了深入的讨论。",
        "他的成绩大约在 90 分左右。",
        "奶奶眯起眼睛说:\"娃儿,你咋个来了哟?\"",
    ]
    for p in paragraphs:
        doc.add_paragraph(p)

    doc.save(str(path))
    print(f"已生成测试文档:{path}")


if __name__ == "__main__":
    out = Path(__file__).parent / "fixtures" / "测试文档.docx"
    make_test_docx(out)
