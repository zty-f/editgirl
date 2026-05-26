"""造一个带专名不一致 + 故意可能改坏的测试 docx,用于 L4/L5 验证。"""
from pathlib import Path
from docx import Document


def make(path: Path):
    doc = Document()
    doc.add_heading("人物故事:张家小镇", level=1)
    # 故意:张伟 出现 3 次,张玮 出现 1 次(同人)
    doc.add_paragraph("张伟从乌鲁木齐回到了家乡,这次他想留下来。")
    doc.add_paragraph("家里人都很高兴,母亲特意做了张伟最爱吃的拌面。")
    doc.add_paragraph("第二天清晨,张玮被鸡叫声唤醒。")  # ← 同人写成"张玮"
    doc.add_paragraph("他走到院子里,看见乌鲁木齐市的天空格外晴朗。")  # ← 乌鲁木齐市 vs 乌鲁木齐
    doc.add_paragraph("回想起在中科院的日子,张伟感慨良多。")
    doc.add_paragraph("毕竟他是 1985 年生人,如今已经四十了。")
    doc.add_paragraph("中国科学院给了他十年的青春。")  # ← 中科院 vs 中国科学院
    doc.add_paragraph("妹妹问他:\"哥,你这次回来还走吗?\"")
    doc.save(str(path))
    print(f"已生成: {path}")


if __name__ == "__main__":
    make(Path(__file__).parent / "fixtures" / "L4测试.docx")
